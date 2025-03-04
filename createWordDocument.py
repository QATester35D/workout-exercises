from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from datetime import datetime
import databaseHelpers
import os

class CreateWorkoutPlan:
    def __init__(self):
        self.document = Document()

    def docWorkOutSelection(self, workoutSelection):
        workoutFocus=("1 - Chest, triceps, abs","2 - Back, biceps, abs","3 - Legs, shoulders, abs","4 - Just legs","5 - Full body workout","6 - HIIT workout","7 - Custom")
        workoutSelectionText=workoutSelection-1
        return(workoutFocus[workoutSelectionText])

    def docCreateHeading(self, workoutFocus):
        document=self.document
        # workoutFocus=("1 - Chest, triceps, abs","2 - Back, biceps, abs","3 - Legs, shoulders, abs","4 - Just legs","5 - Full body workout","6 - HIIT workout","7 - Custom")
        # workoutSelectionText=workOutSelection-1
        document.add_heading('Workout Plan', 0)
        document.add_heading('Brought to you by QATester35D', 1)
        document.add_heading('Muscle Group Focus: '+workoutFocus, 2)
        document.add_heading('This workout plan was created on: '+str(datetime.now()), 3)
        
    def docCreateParagraph(self, workoutFocus):
        document=self.document
        p = document.add_paragraph('\n')
        p = document.add_paragraph('This document is a workout plan created by your selection of the muscle group specified.')
        paragraph_format = p.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('\n')
        p.add_run(workoutFocus).bold = True
        p.add_run('\n')
        p.add_run('These are a random selection of exercises for each muscle group.').italic = True

    def docWriteOutExercises(self, exerciseList):
        document=self.document
        for exercise in exerciseList:
            #Exercise 6 and 7 came from blobs so they need to be decoded.
            # element6=exercise[6].decode('utf-8')
            # lenElement6=len(element6)
            # element6=element6[2:lenElement6-1]
            # element7=exercise[7].decode('utf-8')
            # lenElement7=len(element7)
            # element7=element7[2:lenElement7-1]
            p = document.add_paragraph('\n')
            p.add_run('Exercise ID: '+str(exercise[0])).italic = True
            p.add_run('\n')
            p.add_run('Major Muscle Group: '+exercise[1]).italic = True
            p.add_run('\n')
            p.add_run('Equipment: '+exercise[2]).italic = True
            p.add_run('\n')
            p.add_run('Image URL: '+exercise[3]).italic = True
            p.add_run('\n')
            p.add_run('Exercise Name: '+exercise[4]).italic = True
            p.add_run('\n')
            p.add_run('Targeted Muscle Group: '+exercise[5]).italic = True
            p.add_run('\n')
            p.add_run('Secondary Muscles affected: '+exercise[6].decode('utf-8')).italic = True
            p.add_run('\n')
            p.add_run('Instructions: '+exercise[7].decode('utf-8'))
  
    def docSave(self, fileName):
        self.document.save(fileName)


########################################################################################################
# document = Document()

